"""Agent 引擎: 参考 DeepSeek Harness 的最小工具循环.

- 只读工具立即执行; 所有写操作(文件/日程/邮件/交作业)一律生成"提案",
  由用户在界面上确认后才真正执行(提案 10 分钟过期) —— 与 harness 的权限精神一致.
- 供应商: openai 兼容协议(DeepSeek/Kimi/GLM/通义/Ollama/自定义) + anthropic 协议.
- 文件读写严格限制在用户选择的 workspace 内.
"""
from __future__ import annotations

import json
import re
import time
from datetime import date, datetime, timedelta
from pathlib import Path

from ..config import Config
from ..exceptions import PingheError

MAX_ROUNDS = 8
PROPOSAL_TTL = 600  # 10 分钟
from .. import paths as _paths

SESSIONS_DIR = _paths.data_dir() / "agent_sessions"

# ---------------------------------------------------------------- 权限模式
# readonly        只读: 写工具全部禁用
# confirm         操作前确认(默认): 写操作全部走提案
# workspace_write 工作区写入: workspace 内写文档自动执行, 对外操作仍走提案
# full_access     完全访问: 所有写操作立即执行(切换时前端双重确认)
AGENT_MODES = ("readonly", "confirm", "workspace_write", "full_access")

#: 写工具的许可级别: workspace = 工作区内的写操作; external = 对外/全局操作
_TOOL_LEVEL = {
    "create_docx": "workspace",
    "append_to_docx": "workspace",
    "add_schedule_event": "external",
    "send_email": "external",
    "submit_managebac_task": "external",
}


def _mode_of(cfg: Config) -> str:
    return cfg.agent_mode if cfg.agent_mode in AGENT_MODES else "confirm"


def _now_str() -> str:
    now = datetime.now()
    return f"{now.date()} 周{'一二三四五六日'[now.weekday()]} {now.strftime('%H:%M')}"


def _system_prompt(cfg: Config, workspace: str | None) -> str:
    subjects = "、".join(s["subject"] for s in (cfg.selected_lessons or [])) or "(未选课)"
    mode = _mode_of(cfg)
    lines = [
        "你是 Hello Pinghe! 学习助手, 运行在学生自己的电脑上.",
        f"当前时间: {_now_str()}.",
        f"学生已选科目: {subjects}.",
        f"工作目录(workspace): {workspace or '(未设置)'}",
        f"当前权限模式: {mode}.",
        "规则:",
        "1. 需要课表/DDL/成绩/邮件/日程信息时, 先调用工具查询, 不要编造.",
        "2. 写作业时用 read_docx 查看已有文档, 用 create_docx / append_to_docx 产出草稿.",
    ]
    if mode == "readonly":
        lines.append(
            "3. 当前为只读模式: 一切写工具(create_docx/append_to_docx/"
            "add_schedule_event/send_email/submit_managebac_task)都被禁用, "
            "不要尝试调用; 用户需要写操作时应提示他到 Agent 助手页切换权限模式.")
    elif mode == "confirm":
        lines.append(
            "3. 所有写操作(create_docx/append_to_docx/add_schedule_event/send_email/"
            "submit_managebac_task)都只是提案, 由用户确认后执行, 请在提案前说明你要做什么.")
    elif mode == "workspace_write":
        lines.append(
            "3. 当前为工作区写入模式: create_docx/append_to_docx 会直接执行不必确认;"
            " add_schedule_event/send_email/submit_managebac_task 仍走提案, "
            "请在提案前说明你要做什么.")
    else:
        lines.append(
            "3. 当前为完全访问模式: 所有写操作都会立即执行, 不再有确认弹窗."
            " 请先向用户说明你要做什么再执行, 谨慎操作.")
    lines.append("4. 回答使用简体中文, 简洁直接.")
    if not cfg.send_grades_to_llm:
        lines.append("5. 用户关闭了成绩共享: get_grades 会返回错误, 不要反复尝试.")
    return "\n".join(lines)


# ---------------------------------------------------------------- 工具定义
def _tool(name: str, desc: str, props: dict, required: list | None = None) -> dict:
    return {
        "type": "function",
        "function": {
            "name": name,
            "description": desc,
            "parameters": {
                "type": "object",
                "properties": props,
                "required": required or [],
                "additionalProperties": False,
            },
        },
    }


def build_tools() -> list[dict]:
    return [
        _tool("get_timetable", "查询未来 N 天的个人课表(按学生选课生成)",
              {"days": {"type": "integer", "minimum": 1, "maximum": 14}}),
        _tool("get_ddl", "查询未来 N 天的 ManageBac DDL/作业",
              {"days": {"type": "integer", "minimum": 1, "maximum": 60}}),
        _tool("get_class_tasks", "查询单门课或全部课的作业卡(含已截止)",
              {"class_name": {"type": "string", "maxLength": 80}}),
        _tool("get_grades", "查询各科总评成绩(需用户在设置中允许共享成绩)", {}),
        _tool("list_mail", "列出平和邮箱的邮件",
              {"unseen_only": {"type": "boolean"}, "limit": {"type": "integer", "minimum": 1, "maximum": 50}}),
        _tool("read_mail", "读取一封邮件的正文", {"uid": {"type": "string"}}, ["uid"]),
        _tool("get_schedule", "查询本地日程(YYYY-MM-DD 区间)",
              {"day_from": {"type": "string"}, "day_to": {"type": "string"}},
              ["day_from", "day_to"]),
        _tool("list_workspace", "列出 workspace 下的文件",
              {"subdir": {"type": "string", "maxLength": 200}}),
        _tool("read_docx", "读取 workspace 内 Word 文档的文本", {"path": {"type": "string"}}, ["path"]),
        _tool("read_text_file", "读取 workspace 内文本文件(≤8000字符)",
              {"path": {"type": "string"}}, ["path"]),
        _tool("create_docx", "提案: 在 workspace 新建 Word 文档",
              {"path": {"type": "string"}, "title": {"type": "string"},
               "paragraphs": {"type": "array", "items": {"type": "string"}}},
              ["path", "title", "paragraphs"]),
        _tool("append_to_docx", "提案: 向 workspace 内 Word 文档追加段落",
              {"path": {"type": "string"}, "paragraphs": {"type": "array", "items": {"type": "string"}}},
              ["path", "paragraphs"]),
        _tool("add_schedule_event", "提案: 新增本地日程",
              {"day": {"type": "string"}, "time": {"type": "string"},
               "title": {"type": "string"}, "note": {"type": "string"}},
              ["day", "title"]),
        _tool("search_contacts",
              "在邮箱通讯录里按姓名或邮箱片段搜索联系人。"
              "用户提到人名要发邮件/给谁写信时, 先用本工具把名字变成邮箱地址",
              {"query": {"type": "string", "maxLength": 80}}, ["query"]),
        _tool("send_email", "提案: 用平和邮箱发邮件。to 必须是完整邮箱地址, "
              "如果用户只说了名字, 先调用 search_contacts 查到邮箱",
              {"to": {"type": "string"}, "subject": {"type": "string"}, "body": {"type": "string"}},
              ["to", "subject", "body"]),
        _tool("submit_managebac_task", "提案: 把 workspace 里的文件提交到 ManageBac 作业",
              {"class_id": {"type": "string"}, "task_id": {"type": "string"},
               "file_path": {"type": "string"}},
              ["class_id", "task_id", "file_path"]),
    ]


class AgentEngine:
    def __init__(self, cfg: Config, services):
        self.cfg = cfg
        self.svc = services
        self.tools = build_tools()
        self.history: list[dict] = []
        self.proposals: dict[str, dict] = {}
        self._pid = 0
        self.on_event = None          # 流式回调: bridge 注入, 把增量推给前端
        self.session_id = datetime.now().strftime("%Y%m%d-%H%M%S")

    # ------------------------------------------------------------ 会话持久化
    def save_session(self) -> str | None:
        if not self.history:
            return None
        try:
            SESSIONS_DIR.mkdir(parents=True, exist_ok=True)
            title = next(
                (m["content"][:30] for m in self.history if m["role"] == "user"),
                "会话",
            )
            (SESSIONS_DIR / f"{self.session_id}.json").write_text(
                json.dumps({"id": self.session_id, "title": title,
                            "history": self.history}, ensure_ascii=False),
                encoding="utf-8",
            )
            return self.session_id
        except Exception:  # noqa: BLE001
            return None

    def list_sessions(self) -> list[dict]:
        if not SESSIONS_DIR.exists():
            return []
        out = []
        for p in SESSIONS_DIR.glob("*.json"):
            try:
                data = json.loads(p.read_text(encoding="utf-8"))
                out.append({
                    "id": data.get("id", p.stem),
                    "title": data.get("title", "会话"),
                    "updated": p.stat().st_mtime,
                })
            except Exception:  # noqa: BLE001
                continue
        out.sort(key=lambda s: s["updated"], reverse=True)
        return out[:30]

    def load_session(self, sid: str) -> dict:
        self.save_session()
        data = json.loads((SESSIONS_DIR / f"{sid}.json").read_text(encoding="utf-8"))
        self.history = data.get("history", [])
        self.proposals.clear()
        self.session_id = sid
        return {"session": sid, "title": data.get("title", ""), "history": self.history}

    def new_session(self) -> dict:
        self.save_session()
        self.history = []
        self.proposals.clear()
        self.session_id = datetime.now().strftime("%Y%m%d-%H%M%S")
        return {"session": self.session_id}

    def emit(self, obj: dict) -> None:
        if self.on_event is None:
            return
        try:
            self.on_event(obj)
        except Exception:  # noqa: BLE001
            pass

    # ------------------------------------------------------------ workspace
    def workspace_root(self) -> Path | None:
        if not self.cfg.agent_workspace:
            return None
        return Path(self.cfg.agent_workspace)

    def _remember_workspace(self, path: str) -> None:
        if path and path not in (self.cfg.agent_workspaces or []):
            self.cfg.agent_workspaces = (self.cfg.agent_workspaces or []) + [path]

    def set_workspace(self, path: str) -> dict:
        p = Path(path).expanduser()
        if not p.exists():
            p.mkdir(parents=True, exist_ok=True)
        self.cfg.agent_workspace = str(p)
        self._remember_workspace(str(p))
        return {"workspace": str(p)}

    def new_workspace(self, name: str) -> dict:
        name = re.sub(r'[\\/:*?"<>|]', "_", name or "workspace").strip() or "workspace"
        base = Path.home() / "Documents" / "Hello Pinghe"
        p = base / name
        p.mkdir(parents=True, exist_ok=True)
        self.cfg.agent_workspace = str(p)
        self._remember_workspace(str(p))
        return {"workspace": str(p)}

    def list_workspaces(self) -> list[str]:
        return list(self.cfg.agent_workspaces or [])

    def _resolve(self, rel: str) -> Path:
        root = self.workspace_root()
        if root is None:
            raise PingheError("未设置 workspace, 请先在 Agent 页面选择")
        p = (root / rel).resolve() if rel else root
        if root.resolve() not in p.parents and p != root.resolve():
            raise PingheError(f"路径越出 workspace: {rel}")
        return p

    # ------------------------------------------------------------ 权限模式
    @property
    def mode(self) -> str:
        return _mode_of(self.cfg)

    # ------------------------------------------------------------ 提案
    def _propose(self, title: str, detail: str, fn, level: str = "external") -> dict:
        mode = self.mode
        if mode == "full_access" or (mode == "workspace_write" and level == "workspace"):
            # 高权限模式(开启时已经过双重确认): 直接执行, 不再逐次弹确认
            return {"executed": True, "title": title, "result": fn()}
        self._gc_proposals()
        self._pid += 1
        pid = f"p{int(time.time())}-{self._pid}"
        self.proposals[pid] = {
            "id": pid, "title": title, "detail": detail, "fn": fn,
            "created": time.monotonic(),
        }
        return {"proposal": True, "proposal_id": pid, "title": title, "detail": detail}

    def _gc_proposals(self) -> None:
        now = time.monotonic()
        for pid in [k for k, v in self.proposals.items() if now - v["created"] > PROPOSAL_TTL]:
            self.proposals.pop(pid, None)

    def list_proposals(self) -> list[dict]:
        self._gc_proposals()
        return [
            {"id": v["id"], "title": v["title"], "detail": v["detail"]}
            for v in self.proposals.values()
        ]

    def confirm(self, pid: str) -> dict:
        item = self.proposals.pop(pid, None)
        if not item:
            raise PingheError(f"提案 {pid} 不存在或已过期")
        result = item["fn"]()
        return {"ok": True, "result": result, "title": item["title"]}

    def reject(self, pid: str) -> dict:
        self.proposals.pop(pid, None)
        return {"ok": True}

    # ------------------------------------------------------------ 工具执行
    def _exec_tool(self, name: str, args: dict) -> str:
        try:
            # 只读模式: 一切写工具直接拒绝(LLM 会看到错误并向用户解释)
            if name in _TOOL_LEVEL and self.mode == "readonly":
                raise PingheError(
                    f"当前 Agent 为只读模式, 无法执行写操作({name})。"
                    "请告诉用户: 到 Agent 助手页把权限模式切换为"
                    "「操作前确认/工作区写入/完全访问」后再试。")
            return json.dumps({"ok": True, **self._dispatch(name, args)}, ensure_ascii=False)
        except PingheError as exc:
            return json.dumps({"ok": False, "error": str(exc)}, ensure_ascii=False)
        except Exception as exc:  # noqa: BLE001
            return json.dumps({"ok": False, "error": f"{type(exc).__name__}: {exc}"}, ensure_ascii=False)

    def _dispatch(self, name: str, args: dict):
        if name == "get_timetable":
            days = int(args.get("days") or 1)
            out = []
            for offset in range(days):
                day = date.today() + timedelta(days=offset)
                out.append({"day": str(day), "lessons": self.svc.edupage.personal(day)})
            return {"days": out}

        if name == "get_ddl":
            days = int(args.get("days") or 14)
            return {"deadlines": self.svc.courses.deadlines(days=days)}

        if name == "get_class_tasks":
            wanted = (args.get("class_name") or "").strip().lower()
            tasks = self.svc.courses.all_tasks()
            if wanted:
                tasks = [t for t in tasks if wanted in (t["class_name"] or "").lower()]
            return {"count": len(tasks), "tasks": tasks[:60]}

        if name == "get_grades":
            if not self.cfg.send_grades_to_llm:
                raise PingheError("用户未允许把成绩发给 AI(设置里可开启)")
            return {"grades": self.svc.courses.grades()}

        if name == "list_mail":
            return {"mails": self.svc.mail.list_mail(
                unseen_only=bool(args.get("unseen_only")),
                limit=int(args.get("limit") or 20),
            )}

        if name == "read_mail":
            return self.svc.mail.read(str(args["uid"]))

        if name == "search_contacts":
            return {"contacts": self.svc.mail.contacts_search(
                str(args.get("query") or ""), limit=8)}

        if name == "get_schedule":
            return {"events": self.svc.schedule.list_range(args["day_from"], args["day_to"])}

        if name == "list_workspace":
            root = self.workspace_root()
            if root is None:
                raise PingheError("未设置 workspace")
            base = root / (args.get("subdir") or "")
            if not str(base.resolve()).startswith(str(root.resolve())):
                raise PingheError("路径越出 workspace")
            files = []
            for p in sorted(base.rglob("*")):
                if p.is_file() and ".git" not in p.parts:
                    files.append(str(p.relative_to(root)))
                if len(files) >= 200:
                    break
            return {"workspace": str(root), "files": files}

        if name == "read_docx":
            p = self._resolve(args["path"])
            import docx

            doc = docx.Document(str(p))
            paras = [para.text for para in doc.paragraphs if para.text.strip()]
            return {"path": str(p), "paragraphs": paras[:200]}

        if name == "read_text_file":
            p = self._resolve(args["path"])
            text = p.read_text(encoding="utf-8", errors="replace")
            return {"path": str(p), "text": text[:8000]}

        if name == "create_docx":
            path, title = args["path"], args["title"]
            paras = [str(x) for x in (args.get("paragraphs") or [])]

            def fn():
                p = self._resolve(path)
                if p.suffix.lower() != ".docx":
                    p = p.with_suffix(".docx")
                import docx

                doc = docx.Document()
                doc.add_heading(title, level=1)
                for para in paras:
                    doc.add_paragraph(para)
                doc.save(str(p))
                return f"已创建 {p}"

            return self._propose(f"新建 Word: {path}", f"标题: {title}\n段落数: {len(paras)}", fn,
                                 level="workspace")

        if name == "append_to_docx":
            path = args["path"]
            paras = [str(x) for x in (args.get("paragraphs") or [])]

            def fn():
                p = self._resolve(path)
                import docx

                doc = docx.Document(str(p))
                for para in paras:
                    doc.add_paragraph(para)
                doc.save(str(p))
                return f"已追加 {len(paras)} 段到 {p}"

            return self._propose(f"追加 Word: {path}", f"追加 {len(paras)} 段", fn,
                                 level="workspace")

        if name == "add_schedule_event":
            def fn():
                event_id = self.svc.schedule.add(
                    args["day"], args.get("time") or "", args["title"], args.get("note") or ""
                )
                return f"已添加日程 #{event_id}: {args['title']}"

            return self._propose(
                f"新增日程: {args['title']}",
                f"{args['day']} {args.get('time') or ''} {args.get('note') or ''}", fn
            )

        if name == "send_email":
            def fn():
                self.svc.mail.send(args["to"], args["subject"], args["body"])
                return f"已发送给 {args['to']}: {args['subject']}"

            return self._propose(
                f"发邮件给 {args['to']}", f"主题: {args['subject']}\n---\n{args['body'][:500]}", fn
            )

        if name == "submit_managebac_task":
            def fn():
                p = self._resolve(args["file_path"])
                return self.svc.courses.submit_task(
                    args["class_id"], args["task_id"], str(p)
                )

            return self._propose(
                f"提交作业 {args['task_id']}",
                f"课程 {args['class_id']} ← {args['file_path']}", fn
            )

        raise PingheError(f"未知工具: {name}")

    # ------------------------------------------------------------ LLM 调用
    def _active(self) -> tuple[dict, str]:
        provider = self.cfg.active_provider()
        model = self.cfg.agent_model or (provider.get("models") or [""])[0]
        return provider, model

    def _call_llm(self, messages: list[dict], on_delta=None):
        provider, model = self._active()
        if provider.get("protocol") == "anthropic":
            return self._call_anthropic(provider, model, messages, on_delta)
        return self._call_openai(provider, model, messages, on_delta)

    def _call_openai(self, provider: dict, model: str,
                     messages: list[dict], on_delta=None):
        from openai import OpenAI

        client = OpenAI(
            api_key=provider.get("api_key") or "EMPTY",
            base_url=provider.get("base_url") or None,
        )
        content_parts: list[str] = []
        tool_calls: dict[int, dict] = {}

        stream = client.chat.completions.create(
            model=model,
            messages=messages,
            tools=self.tools,
            max_tokens=4000,
            stream=True,
        )
        for chunk in stream:
            if not chunk.choices:
                continue
            delta = chunk.choices[0].delta
            if delta is None:
                continue
            if delta.content:
                content_parts.append(delta.content)
                if on_delta:
                    on_delta(delta.content)
            if delta.tool_calls:
                for tc in delta.tool_calls:
                    idx = tc.index if tc.index is not None else len(tool_calls)
                    slot = tool_calls.setdefault(idx, {"id": "", "name": "", "arguments": ""})
                    if tc.id:
                        slot["id"] = tc.id
                    if tc.function:
                        if tc.function.name:
                            slot["name"] += tc.function.name
                        if tc.function.arguments:
                            slot["arguments"] += tc.function.arguments

        calls = [tool_calls[i] for i in sorted(tool_calls)]
        return {"content": "".join(content_parts), "tool_calls": calls}

    def _call_anthropic(self, provider: dict, model: str,
                        messages: list[dict], on_delta=None):
        import anthropic

        base_url = (provider.get("base_url") or "").strip().rstrip("/") or None
        client = anthropic.Anthropic(
            api_key=provider.get("api_key"),
            base_url=base_url,
        )
        # 内部历史是 openai 风格, 转成 anthropic 风格
        conv, pending_tool_results = [], []
        for m in messages:
            role, content = m["role"], m.get("content") or ""
            if m["role"] == "assistant" and m.get("tool_calls"):
                blocks = []
                if content:
                    blocks.append({"type": "text", "text": content})
                for tc in m["tool_calls"]:
                    blocks.append({
                        "type": "tool_use", "id": tc["id"], "name": tc["name"],
                        "input": json.loads(tc["arguments"] or "{}"),
                    })
                conv.append({"role": "assistant", "content": blocks})
                continue
            if m["role"] == "tool":
                pending_tool_results.append({
                    "type": "tool_result", "tool_use_id": m["tool_call_id"],
                    "content": str(content),
                })
                continue
            if pending_tool_results:
                conv.append({"role": "user", "content": pending_tool_results})
                pending_tool_results = []
            conv.append({"role": role, "content": content or "(空)"})
        if pending_tool_results:
            conv.append({"role": "user", "content": pending_tool_results})

        tools = [
            {
                "name": t["function"]["name"],
                "description": t["function"]["description"],
                "input_schema": t["function"]["parameters"],
            }
            for t in self.tools
        ]
        content_parts: list[str] = []
        with client.messages.stream(
            model=model, max_tokens=4000,
            system=_system_prompt(self.cfg, self.cfg.agent_workspace),
            messages=conv, tools=tools,
        ) as stream:
            for text in stream.text_stream:
                content_parts.append(text)
                if on_delta:
                    on_delta(text)
            final = stream.get_final_message()

        content = "".join(content_parts)
        calls = [
            {
                "id": block.id, "name": block.name,
                "arguments": json.dumps(block.input, ensure_ascii=False),
            }
            for block in final.content if block.type == "tool_use"
        ]
        return {"content": content, "tool_calls": calls}

    # ------------------------------------------------------------ 主循环
    def chat(self, message: str) -> dict:
        if not self.history:
            self.history.append({
                "role": "system", "content": _system_prompt(self.cfg, self.cfg.agent_workspace)
            })
        self.history.append({"role": "user", "content": message})

        reply = ""
        for _ in range(MAX_ROUNDS):
            resp = self._call_llm(self.history, on_delta=lambda t: self.emit({"type": "delta", "text": t}))
            calls = resp["tool_calls"]
            if not calls:
                reply = resp["content"]
                self.history.append({"role": "assistant", "content": reply})
                self.save_session()
                return {"ok": True, "reply": reply}

            assistant = {"role": "assistant", "content": resp["content"],
                         "tool_calls": calls}
            self.history.append(assistant)
            for call in calls:
                try:
                    args = json.loads(call["arguments"] or "{}")
                except json.JSONDecodeError:
                    args = {}
                self.emit({"type": "tool", "name": call["name"]})
                result = self._exec_tool(call["name"], args)
                self.emit({"type": "tool_result", "name": call["name"], "preview": result[:240]})
                try:
                    parsed = json.loads(result)
                    if parsed.get("proposal"):
                        self.emit({"type": "proposal"})
                except json.JSONDecodeError:
                    pass
                self.history.append({
                    "role": "tool", "tool_call_id": call["id"],
                    "name": call["name"], "content": result,
                })

        self.save_session()
        return {"ok": False, "error": f"达到最大工具轮数({MAX_ROUNDS}), 已暂停", "reply": reply}

    def reset(self) -> None:
        self.new_session()


def detect_ai_environment() -> dict:
    """粗略检测硬件, 给出 Ollama 模型或 API 建议. (Windows/macOS/Linux 兼容)"""
    import os
    import platform

    system = platform.system()
    ram_gb = 0.0
    gpu = ""
    if system == "Windows":
        try:
            import ctypes


            class MEMORYSTATUSEX(ctypes.Structure):
                _fields_ = [
                    ("dwLength", ctypes.c_ulong),
                    ("dwMemoryLoad", ctypes.c_ulong),
                    ("ullTotalPhys", ctypes.c_ulonglong),
                    ("ullAvailPhys", ctypes.c_ulonglong),
                    ("ullTotalPageFile", ctypes.c_ulonglong),
                    ("ullAvailPageFile", ctypes.c_ulonglong),
                    ("ullTotalVirtual", ctypes.c_ulonglong),
                    ("ullAvailVirtual", ctypes.c_ulonglong),
                    ("ullAvailExtendedVirtual", ctypes.c_ulonglong),
                ]

            stat = MEMORYSTATUSEX()
            stat.dwLength = ctypes.sizeof(MEMORYSTATUSEX)
            ctypes.windll.kernel32.GlobalMemoryStatusEx(ctypes.byref(stat))
            ram_gb = stat.ullTotalPhys / (1024 ** 3)
        except Exception:  # noqa: BLE001
            pass
    else:
        try:
            # POSIX(macOS/Linux): 页大小 × 物理页数
            page = os.sysconf("SC_PAGE_SIZE")
            phys = os.sysconf("SC_PHYS_PAGES")
            ram_gb = page * phys / (1024 ** 3)
        except Exception:  # noqa: BLE001
            pass
    if system == "Windows":
        try:
            import subprocess

            out = subprocess.run(
                ["wmic", "path", "win32_VideoController", "get", "name"],
                capture_output=True, text=True, timeout=10,
            ).stdout
            candidates = [
                line.strip() for line in out.splitlines()
                if line.strip() and "VideoController" not in line and "Name" not in line
            ]
            # 优先真实显卡, 过滤虚拟显示适配器(OrayIdd/向日葵 等)
            def rank(name: str) -> int:
                upper = name.upper()
                if any(k in upper for k in ("NVIDIA", "GEFORCE", "RADEON", "AMD")):
                    return 0
                if "INTEL" in upper and "IDD" not in upper and "VIRTUAL" not in upper:
                    return 1
                if "IDD" in upper or "VIRTUAL" in upper or "MIRROR" in upper:
                    return 9
                return 5
            if candidates:
                gpu = sorted(candidates, key=rank)[0]
        except Exception:  # noqa: BLE001
            pass
    elif system == "Darwin":
        try:
            import subprocess

            out = subprocess.run(
                ["system_profiler", "SPDisplaysDataType"],
                capture_output=True, text=True, timeout=15,
            ).stdout
            m = re.search(r"Chipset Model:\s*(.+)", out)
            gpu = m.group(1).strip() if m else ""
        except Exception:  # noqa: BLE001
            pass

    has_nvidia = "NVIDIA" in (gpu or "").upper()
    if ram_gb >= 16 or has_nvidia:
        model = "qwen2.5:14b" if (ram_gb >= 24 or has_nvidia) else "qwen2.5:7b"
        advice = "local"
    else:
        model = "qwen2.5:3b"
        advice = "api"
    return {
        "platform": platform.platform(),
        "cpu": os.cpu_count(),
        "ram_gb": round(ram_gb, 1),
        "gpu": gpu or "(未检测到)",
        "advice": advice,          # local = 推荐本地模型, api = 建议 API
        "recommended_model": model,
    }
